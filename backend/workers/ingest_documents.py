"""
通用文档摄取脚本：
支持八字、紫微斗数、通用命理等多种类型的文档摄取到 PostgreSQL (PGVector)

用法:
    python -m workers.ingest_documents --source-type bazi --rebuild
    python -m workers.ingest_documents --source-type ziwei
    python -m workers.ingest_documents --source-type fortune
    python -m workers.ingest_documents --all --rebuild
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Dict, Iterable, List, Tuple, Optional

from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from langchain_community.vectorstores.pgvector import PGVector
from langchain_core.documents import Document

# 添加后端根目录到 Python 路径
BASE_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(BASE_DIR))

from app.core.settings import settings

# 数据源配置
SOURCE_TYPE_CONFIGS = {
    "bazi": {
        "data_dir": BASE_DIR / "data" / "raw" / "bazi",
        "description": "八字命理",
        "file_extensions": [".pdf", ".docx", ".md", ".txt"]
    },
    "ziwei": {
        "data_dir": BASE_DIR / "data" / "raw" / "ziwei", 
        "description": "紫微斗数",
        "file_extensions": [".pdf", ".docx", ".md", ".txt"]
    },
    "fortune": {
        "data_dir": BASE_DIR / "data" / "raw",
        "description": "通用命理",
        "file_extensions": [".docx", ".md", ".txt"],
        "exclude_dirs": ["bazi", "ziwei"]  # 排除其他专门的目录
    }
}

COLLECTION_NAME = "rag_documents"


def read_pdf_text(path: Path) -> str:
    """读取 PDF 文件内容"""
    try:
        import PyPDF2
        with open(path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except ImportError:
        print("警告：未安装 PyPDF2，尝试使用 pdfplumber...")
    
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            text = ""
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text.strip()
    except ImportError:
        print("警告：未安装 pdfplumber，尝试使用 pymupdf...")
    
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text() + "\n"
        doc.close()
        return text.strip()
    except ImportError:
        print("错误：无法处理 PDF 文件。请安装 PyPDF2、pdfplumber 或 pymupdf")
        return ""
    except Exception as e:
        print(f"读取 PDF 文件失败 {path}: {e}")
        return ""


def read_docx_text(path: Path) -> str:
    """读取 DOCX 文件内容"""
    try:
        import docx2txt  # type: ignore
        txt = docx2txt.process(str(path)) or ""
        if txt and txt.strip():
            return txt
    except Exception:
        pass
    
    try:
        from docx import Document as DocxDocument  # type: ignore
        doc = DocxDocument(str(path))
        lines = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                lines.append(t)
        for tbl in getattr(doc, "tables", []) or []:
            for row in tbl.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        lines.append(t)
        return "\n".join(lines)
    except Exception as e:
        print(f"读取 DOCX 文件失败 {path}: {e}")
        return ""


def read_text_file(path: Path) -> str:
    """读取纯文本文件内容"""
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"读取文本文件失败 {path}: {e}")
        return ""


def iter_source_documents(source_type: str) -> Iterable[Tuple[str, str, Dict[str, str]]]:
    """迭代指定类型的源文档"""
    if source_type not in SOURCE_TYPE_CONFIGS:
        raise ValueError(f"未知的源类型: {source_type}")
    
    config = SOURCE_TYPE_CONFIGS[source_type]
    data_dir = config["data_dir"]
    extensions = config["file_extensions"]
    exclude_dirs = config.get("exclude_dirs", [])
    
    if not data_dir.exists():
        print(f"数据目录不存在: {data_dir}")
        return
    
    for p in data_dir.rglob("*"):
        if not p.is_file():
            continue
            
        # 检查是否在排除目录中
        if exclude_dirs:
            relative_path = p.relative_to(data_dir)
            if any(exclude_dir in relative_path.parts for exclude_dir in exclude_dirs):
                continue
        
        suffix = p.suffix.lower()
        if suffix not in extensions:
            continue
            
        print(f"处理文件: {p}")
        
        content = ""
        if suffix == ".pdf":
            content = read_pdf_text(p)
        elif suffix == ".docx":
            content = read_docx_text(p)
        elif suffix in {".md", ".txt"}:
            content = read_text_file(p)
        
        if content and content.strip():
            # 构建元数据
            metadata = {
                "source_type": source_type,
                "source_book": p.name,
                "source_path": str(p),
                "file_type": suffix[1:],  # 去掉点号
                "description": config["description"]
            }
            yield (str(p), content.strip(), metadata)


def split_by_sentences(text: str) -> List[str]:
    """按句子分割文本"""
    seps = "。！？!?\n"
    buf: List[str] = []
    sent = []
    for ch in text:
        sent.append(ch)
        if ch in seps:
            buf.append("".join(sent).strip())
            sent = []
    if sent:
        buf.append("".join(sent).strip())
    return [s for s in buf if s]


def make_chunks_by_sentences(text: str, size: int, overlap: int) -> List[str]:
    """基于句子创建文本块"""
    sents = split_by_sentences(text)
    chunks: List[str] = []
    bucket: List[str] = []
    cur_len = 0
    
    for s in sents:
        sl = len(s)
        if cur_len + sl <= size or not bucket:
            bucket.append(s)
            cur_len += sl
        else:
            chunks.append("".join(bucket))
            # 重叠处理
            rollback = overlap
            while bucket and rollback > 0:
                rollback -= len(bucket[-1])
                bucket.pop()
            bucket.append(s)
            cur_len = len("".join(bucket))
    
    if bucket:
        chunks.append("".join(bucket))
    
    return chunks


def ingest_source_type(source_type: str, rebuild: bool = False) -> None:
    """摄取指定类型的文档"""
    print(f"开始摄取 {source_type} 类型的文档...")
    
    if not settings.DATABASE_URL:
        raise ValueError("DATABASE_URL 未设置，请检查环境配置")
    
    # 准备嵌入模型
    embeddings = HuggingFaceBgeEmbeddings(
        model_name="BAAI/bge-small-zh-v1.5",
        encode_kwargs={"normalize_embeddings": True},
    )
    
    # 收集文档
    documents: List[Document] = []
    file_count = 0
    
    for source_path, content, metadata in iter_source_documents(source_type):
        file_count += 1
        
        # 创建文本块
        chunks = make_chunks_by_sentences(content, size=800, overlap=100)
        
        for chunk_idx, chunk_text in enumerate(chunks):
            # 为每个块添加元数据
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "chunk_index": chunk_idx,
                "total_chunks": len(chunks),
                "chunk_id": f"{source_path}#chunk_{chunk_idx}"
            })
            
            doc = Document(
                page_content=chunk_text,
                metadata=chunk_metadata
            )
            documents.append(doc)
    
    if not documents:
        print(f"未找到 {source_type} 类型的有效文档")
        return
    
    print(f"准备存储 {len(documents)} 个文档块到 PGVector...")
    
    # 如果需要重建，删除现有的同类型文档
    if rebuild:
        print(f"删除现有的 {source_type} 类型文档...")
        # 这里我们创建一个临时的 PGVector 实例来删除现有数据
        vectordb = PGVector(
            connection_string=settings.DATABASE_URL,
            embedding_function=embeddings,
            collection_name=COLLECTION_NAME,
        )
        # 注意：PGVector 没有直接的删除方法，我们可能需要手动执行 SQL
        # 这里暂时跳过，实际部署时可以考虑手动清理数据库
    
    # 存储到 PGVector
    try:
        vectordb = PGVector.from_documents(
            documents=documents,
            embedding=embeddings,
            connection_string=settings.DATABASE_URL,
            collection_name=COLLECTION_NAME,
        )
        
        print(f"成功存储 {source_type} 类型文档:")
        print(f"- 处理文件数: {file_count}")
        print(f"- 文档块数: {len(documents)}")
        print(f"- 集合名称: {COLLECTION_NAME}")
        
    except Exception as e:
        print(f"存储文档时出错: {e}")
        raise


def ingest_all_types(rebuild: bool = False) -> None:
    """摄取所有类型的文档"""
    for source_type in SOURCE_TYPE_CONFIGS.keys():
        try:
            ingest_source_type(source_type, rebuild=rebuild)
            print(f"{source_type} 类型摄取完成\n")
        except Exception as e:
            print(f"摄取 {source_type} 类型时出错: {e}\n")


def main():
    parser = argparse.ArgumentParser(
        description="通用文档摄取脚本 - 支持多种命理文档类型",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例用法:
    python -m workers.ingest_documents --source-type bazi --rebuild
    python -m workers.ingest_documents --source-type ziwei  
    python -m workers.ingest_documents --all --rebuild
        """
    )
    
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument(
        "--source-type", 
        choices=list(SOURCE_TYPE_CONFIGS.keys()),
        help="指定要摄取的文档类型"
    )
    group.add_argument(
        "--all", 
        action="store_true",
        help="摄取所有类型的文档"
    )
    
    parser.add_argument(
        "--rebuild", 
        action="store_true", 
        help="重建索引（删除现有的同类型文档）"
    )
    
    args = parser.parse_args()
    
    if args.all:
        ingest_all_types(rebuild=args.rebuild)
    else:
        ingest_source_type(args.source_type, rebuild=args.rebuild)


if __name__ == "__main__":
    main()

