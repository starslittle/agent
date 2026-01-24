"""
命理库入库脚本（迁移版）：

- 遍历 data/raw 下的 .docx/.md/.txt
- 生成父子分块，并写入 ./storage/chroma/fortune
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable, List, Tuple

from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document


BASE_DIR = Path(__file__).resolve().parents[2]
DATA_DIR = BASE_DIR / "data" / "raw"
DB_DIR = BASE_DIR / "storage" / "chroma" / "fortune"
SUMMARY_COLLECTION = "fortune_summary"
PASSAGE_COLLECTION = "fortune"


def read_docx_text(path: Path) -> str:
    try:
        import docx2txt  # type: ignore
        txt = docx2txt.process(str(path)) or ""
        if txt and txt.strip():
            return txt
    except Exception:
        pass
    try:
        from docx import Document  # type: ignore
        doc = Document(str(path))
        lines = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                lines.append(t)
        for tbl in getattr(doc, "tables", []) or []:
            for row in tbl.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        lines.append(t)
        return "\n".join(lines)
    except Exception:
        return ""


def read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def iter_raw_texts() -> Iterable[Tuple[str, str]]:
    for p in DATA_DIR.rglob("*"):
        if not p.is_file():
            continue
        suffix = p.suffix.lower()
        content = ""
        if suffix == ".docx":
            content = read_docx_text(p)
        elif suffix in {".md", ".txt"}:
            content = read_text_file(p)
        if content and content.strip():
            yield (str(p), content.strip())


def split_by_sentences(text: str) -> List[str]:
    seps = "。！？!?\n"
    buf: List[str] = []
    sent = []
    for ch in text:
        sent.append(ch)
        if ch in seps:
            buf.append("".join(sent).strip())
            sent = []
    if sent:
        buf.append("".join(sent).strip())
    return [s for s in buf if s]


def make_chunks_by_sentences(text: str, size: int, overlap: int) -> List[str]:
    sents = split_by_sentences(text)
    chunks: List[str] = []
    bucket: List[str] = []
    cur_len = 0
    for s in sents:
        sl = len(s)
        if cur_len + sl <= size or not bucket:
            bucket.append(s)
            cur_len += sl
        else:
            chunks.append("".join(bucket))
            rollback = overlap
            while bucket and rollback > 0:
                rollback -= len(bucket[-1])
                bucket.pop()
            bucket.append(s)
            cur_len = len("".join(bucket))
    if bucket:
        chunks.append("".join(bucket))
    return chunks


def build_summary(text: str, limit: int = 400) -> str:
    sents = split_by_sentences(text)
    acc = []
    total = 0
    for s in sents:
        if total + len(s) > limit:
            break
        acc.append(s)
        total += len(s)
    return "".join(acc) if acc else text[:limit]


def ingest(rebuild: bool = False) -> None:
    if rebuild and DB_DIR.exists():
        import shutil
        shutil.rmtree(DB_DIR, ignore_errors=True)
    DB_DIR.mkdir(parents=True, exist_ok=True)

    summary_docs: List[Document] = []
    passage_docs: List[Document] = []
    file_count = 0

    for path, content in iter_raw_texts():
        file_count += 1
        parents = make_chunks_by_sentences(content, size=1800, overlap=200)
        for idx, parent_txt in enumerate(parents):
            parent_id = f"{path}#P{idx}"
            summary_docs.append(
                Document(page_content=build_summary(parent_txt, 400), metadata={"parent_id": parent_id, "source": path})
            )
            children = make_chunks_by_sentences(parent_txt, size=700, overlap=120)
            for j, ch_txt in enumerate(children):
                passage_docs.append(
                    Document(page_content=ch_txt, metadata={"parent_id": parent_id, "source": path, "child_idx": j})
                )

    if not passage_docs:
        print(f"未在 {DATA_DIR} 发现可用文档（.docx/.md/.txt），已跳过。")
        return

    embeddings = HuggingFaceBgeEmbeddings(
        model_name="BAAI/bge-small-zh-v1.5",
        encode_kwargs={"normalize_embeddings": True},
    )
    Chroma.from_documents(
        documents=summary_docs,
        embedding=embeddings,
        persist_directory=str(DB_DIR),
        collection_name=SUMMARY_COLLECTION,
    )
    Chroma.from_documents(
        documents=passage_docs,
        embedding=embeddings,
        persist_directory=str(DB_DIR),
        collection_name=PASSAGE_COLLECTION,
    )
    print(
        f"已处理文件数: {file_count}，摘要条目: {len(summary_docs)}，子块条目: {len(passage_docs)} → {DB_DIR}"
    )


def main():
    parser = argparse.ArgumentParser(description="Ingest fortune knowledge base")
    parser.add_argument("--rebuild", action="store_true", help="重建索引（删除旧库）")
    args = parser.parse_args()
    ingest(rebuild=args.rebuild)


if __name__ == "__main__":
    main()


